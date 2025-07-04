from flask import Flask, render_template, request, send_file, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
import os
from datetime import datetime

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

GENERATED_DIR = 'generated_docs'
os.makedirs(GENERATED_DIR, exist_ok=True)

class Petition(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    court = db.Column(db.String(100), nullable=False)
    location = db.Column(db.String(100), nullable=False)
    mp_no = db.Column(db.String(50), nullable=False)
    cc_no = db.Column(db.String(50), nullable=False)
    petitioner = db.Column(db.String(200), nullable=False)
    respondent = db.Column(db.String(200), nullable=False)
    reason = db.Column(db.Text, nullable=False)
    place = db.Column(db.String(100), nullable=False)
    date = db.Column(db.String(20), nullable=False)
    generated_filename = db.Column(db.String(255), unique=True, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):
        return f"Petition('{self.petitioner}' vs '{self.respondent}' - MP No: {self.mp_no})"

with app.app_context():
    db.create_all()

def add_centered_heading(doc, text, size=12, bold=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_right_aligned_text(doc, text, size=12, bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)

@app.route('/', methods=['GET', 'POST'])
def form_317():
    if request.method == 'POST':
        court = request.form['court']
        location = request.form['location']
        mp_no = request.form['mp_no']
        cc_no = request.form['cc_no']
        petitioner = request.form['petitioner']
        respondent = request.form['respondent']
        reason = request.form['reason']
        place = request.form['place']
        date = request.form['date']

        filename = f"petition_317_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join(GENERATED_DIR, filename)

        doc = Document()
        add_centered_heading(doc, f"IN THE COURT OF {court.upper()}", size=14)
        add_centered_heading(doc, location.upper(), size=12)

        current_year = datetime.now().year
        add_centered_heading(doc, f"M.P. No. {mp_no} of {current_year}")
        add_centered_heading(doc, "In")
        add_centered_heading(doc, f"C.C. No. {cc_no} of {current_year}")
        doc.add_paragraph()

        section = doc.sections[0]
        tab_stop_pos = Inches(6)

        p_petitioner = doc.add_paragraph()
        p_petitioner.add_run(f"{petitioner.upper()}")
        p_petitioner.paragraph_format.tab_stops.add_tab_stop(tab_stop_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        p_petitioner.add_run("\t… PETITIONER / ACCUSED.")

        add_centered_heading(doc, "VS")

        p_respondent = doc.add_paragraph()
        p_respondent.add_run(f"{respondent.upper()}")
        p_respondent.paragraph_format.tab_stops.add_tab_stop(tab_stop_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        p_respondent.add_run("\t… RESPONDENT / COMPLAINANT.")
        doc.add_paragraph()

        add_centered_heading(doc, "PETITION FILED UNDER SECTION 317 Cr.P.C.", size=12)

        doc.add_paragraph("The Petitioner/Accused above named most respectfully states as follows:\n")
        doc.add_paragraph("1. That the Petitioner states that the above case is posted today for further proceedings.")
        doc.add_paragraph(f"2. The Petitioner further states that the petitioner is unable to appear before this Hon’ble Court today due to {reason}.")
        doc.add_paragraph("3. That the absence of the Petitioner before this Hon’ble Court is neither willful nor wanton but purely due to the reason stated above.")

        doc.add_paragraph(
            "\nIn these circumstances, it is prayed that this Hon’ble Court may be pleased to dispense with "
            "the personal appearance of the accused for today only and thus render justice.\n"
        )
        doc.add_paragraph()
        doc.add_paragraph(f"Place : {place}")
        doc.add_paragraph(f"Date  : {date}")
        add_right_aligned_text(doc, "Counsel for Petitioner / Accused")

        doc.save(filepath)

        new_petition = Petition(
            court=court,
            location=location,
            mp_no=mp_no,
            cc_no=cc_no,
            petitioner=petitioner,
            respondent=respondent,
            reason=reason,
            place=place,
            date=date,
            generated_filename=filename
        )
        db.session.add(new_petition)
        db.session.commit()

        return send_file(filepath, as_attachment=True)

    today = datetime.now().strftime('%d-%m-%Y')
    return render_template('form_317.html', today=today)

@app.route('/petitions')
def view_petitions():
    petitions = Petition.query.order_by(Petition.timestamp.desc()).all()
    return render_template('view_petitions.html', petitions=petitions)

@app.route('/download/<filename>')
def download_file(filename):
    filepath = os.path.join(GENERATED_DIR, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        return "File not found.", 404

if __name__ == '__main__':
    app.run(debug=True)
