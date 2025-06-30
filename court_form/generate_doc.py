from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_affidavit_doc(data):
    doc = Document()

    # --- PAGE 1: AFFIDAVIT ---
    p = doc.add_paragraph(f"IN THE COURT OF THE {data['court'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(f"AT {data['location'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph(f"I.A. No: {data['ia_no']}    OP/OS No: {data['op_no']}")

    para = doc.add_paragraph("BETWEEN:")
    para.runs[0].bold = True
    doc.add_paragraph(f"{data['petitioner']}\n(Petitioner)")

    para = doc.add_paragraph("AND")
    para.runs[0].bold = True
    doc.add_paragraph(f"{data['respondent']}\n(Respondent)")

    heading = doc.add_paragraph("AFFIDAVIT")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True

    doc.add_paragraph(
        f"I, {data['petitioner']}, S/D/O {data['petitioner_son_of']}, aged {data['age']} years, "
        f"Occupation: {data['occupation']}, Resident of {data['address']}, "
        "do hereby solemnly affirm and state as follows:"
    )

    doc.add_paragraph("1. I am the Petitioner in the above matter and am well acquainted with the facts of the case.")

    doc.add_paragraph(
        "2. The Respondent has been avoiding service of summons by keeping out of the way and the summons/notice could not be served in the usual course."
    )

    doc.add_paragraph(
        "Hence, it is prayed that this Hon'ble Court may be pleased to pass an order for substitute service of summons "
        "on the Defendant/Respondent by publication in any local daily newspaper or any other newspaper and to pass such "
        "other orders as deemed fit and proper in the circumstances of the case."
    )

    doc.add_paragraph(f"Place: {data['place']}".ljust(50) + "Deponent Signature: __________________________")
    doc.add_paragraph(f"Date: {data['date']}")
    doc.add_paragraph("Sworn and signed before me on this ____ day of ____________, 20___")
    doc.add_paragraph("Advocate Signature: __________________________")

    doc.add_page_break()

    # --- PAGE 2: PETITION UNDER CPC ---
    p = doc.add_paragraph(f"IN THE COURT OF THE {data['court'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(f"AT {data['location'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph(f"I.A. No: {data['ia_no']}    OP/OS No: {data['op_no']}")
    doc.add_paragraph("BETWEEN:")
    doc.add_paragraph(f"{data['petitioner']}\n(Petitioner)")
    doc.add_paragraph("AND")
    doc.add_paragraph(f"{data['respondent']}\n(Respondent)")

    # Centered, bold heading
    p = doc.add_paragraph("Petition filed under Order V Rule 20 C.P.C.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    # Full petition body
    doc.add_paragraph(
        "The Petitioner submits that despite several efforts, the summons/notice to the Respondent could not be served in the normal course. "
        "The Respondent is deliberately avoiding service and is not available at their usual place of residence or work. "
        "It is, therefore, necessary to seek permission of this Hon’ble Court to effect service of summons through substituted means as per Order V Rule 20 of the Code of Civil Procedure, 1908."
    )

    doc.add_paragraph(
        "The Petitioner, therefore, prays that this Hon’ble Court may be pleased to grant permission to serve summons on the Respondent "
        "by way of publication in any widely circulated daily newspaper, either in English or the local language, and to pass such other "
        "order(s) as may be deemed fit and proper in the circumstances of the case."
    )

    doc.add_paragraph("Filed By:")
    doc.add_paragraph(f"{data['petitioner']}\n(Petitioner)")
    doc.add_paragraph("Counsel for the Petitioner")

    doc.add_page_break()

    # --- PAGE 3: FILING DETAILS ---
    p = doc.add_paragraph(f"IN THE COURT OF THE {data['court'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(f"AT {data['location'].upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph(f"I.A. No: {data['ia_no']}    OP/OS No: {data['op_no']}")
    doc.add_paragraph("BETWEEN:")
    doc.add_paragraph(f"{data['petitioner']}\n(Petitioner)")
    doc.add_paragraph("AND")
    doc.add_paragraph(f"{data['respondent']}\n(Respondent)")

    doc.add_paragraph(f"Filed on: {data['date']}")
    doc.add_paragraph(f"Filed by: {data['petitioner']}")
    doc.add_paragraph(f"Address for Service: {data['address']}")

    # Centered and bold again
    p = doc.add_paragraph("Petition Filed under Order V Rule 20 C.P.C.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("Respondents/Defendants")
    doc.add_paragraph("TELANGANA ADVOCATES' MUTUALLY AIDED")
    doc.add_paragraph("CO-OPERATIVE SOCIETY LIMITED, HYDERABAD, T.S.")

    output_path = f"Affidavit_{data['petitioner'].replace(' ', '_')}.docx"
    doc.save(output_path)
    return output_path
