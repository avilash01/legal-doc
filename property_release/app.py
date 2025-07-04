from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from io import BytesIO

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def release_form():
    if request.method == 'POST':
        data = request.form.to_dict()

        # Create a Word document
        doc = Document()

        # Add centered, bold, black heading
        heading = doc.add_paragraph()
        run = heading.add_run("PERSONAL PROPERTY RELEASE")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Introductory paragraph
        doc.add_paragraph(
            f'THIS PERSONAL PROPERTY RELEASE, made this {data["day"]} day of {data["month"]}, {data["year"]}, '
            f'between {data["grantor_name"]} ("Grantor(s)") who currently resides at {data["grantor_address"]} '
            f'and {data["grantee_name"]} ("Grantee") whose address is {data["grantee_address"]}.'
        )

        # Terms of release
        doc.add_paragraph(
            f'WHEREAS, Grantor agrees to vacate said Property no later than {data["vacancy_date"]} ("Vacancy Date") '
            f'and to leave the property in broom-swept condition free of interior and exterior trash, debris or damage '
            f'and to remove all personal property. Grantor agrees that Grantee is authorized to access and dispose of any '
            f'personal property remaining after the Vacancy Date. This Release shall be enforceable by Grantee and its '
            f'successors, agents, or assigns. This document may be relied upon as proof of Grantor’s consent for such removal or disposal.'
        )

        # Liability clause
        doc.add_paragraph(
            f'In consideration of Grantee’s acceptance of a mortgage release/deed-in-lieu of foreclosure of the Property, '
            f'Grantor(s) hereby releases and holds harmless Grantee and its servicers, representatives, agents, attorneys, '
            f'Officers, Directors, employees, successors and assigns from any claim or liability, loss, cost, or expense, '
            f'including reasonable attorney’s fees, for any and all personal property left in the Property after the agreed upon Vacancy Date.'
        )

        # Agreement statement
        doc.add_paragraph(
            f'THE UNDERSIGNED GRANTOR(S) HAS READ THE FOREGOING PERSONAL PROPERTY RELEASE AND UNDERSTANDS IT, AND HAS HAD '
            f'THE OPPORTUNITY TO CONSULT WITH COUNSEL AND AGREES THAT ALL DOUBTS AND AMBIGUITIES IN CONNECTION WITH THIS '
            f'PERSONAL PROPERTY RELEASE SHALL BE CONSTRUED IN FAVOR OF THE RELEASED PARTY.'
        )

        # Signature section
        doc.add_paragraph("\n\nBy:\n\nGrantor")
        doc.add_paragraph("\n\nBy:\n\nGrantee")
        doc.add_paragraph("Witness:")
        doc.add_paragraph("Witness")  # ← You wanted this word, not the name

        # Return the document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name="personal_property_release.docx")

    return render_template("form_property_release.html")


if __name__ == "__main__":
    app.run(debug=True)
