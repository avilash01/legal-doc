from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_property_release_doc(data):
    doc = Document()

    # Title
    title = doc.add_paragraph("PERSONAL PROPERTY RELEASE")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    # Body
    doc.add_paragraph(
        f"This PERSONAL PROPERTY RELEASE, made this {data['day']} day of {data['month']}, {data['year']}, "
        f"between {data['grantor']} ('Grantor') who currently resides at {data['grantor_address']} ('Property') "
        f"and {data['grantee']} ('Grantee') whose address is {data['grantee_address']}."
    )

    doc.add_paragraph(
        f"WHEREAS, Grantor agrees to vacate said Property no later than {data['vacancy_date']} ('Vacancy Date') "
        "and to leave the property in broom-swept condition free of interior and exterior trash, debris or damage and to remove all personal property. "
        "Grantor agrees that Grantee is authorized to access and dispose of any personal property remaining after the Vacancy Date. "
        "This Release shall be enforceable by Grantee and its successors, agents, or assigns. "
        "This document may be relied upon as proof of Grantor's consent for such removal or disposal."
    )

    doc.add_paragraph(
        "THE UNDERSIGNED GRANTOR(S) HAS READ THE FOREGOING PERSONAL PROPERTY RELEASE AND UNDERSTANDS IT, "
        "HAS HAD THE OPPORTUNITY TO CONSULT WITH COUNSEL, AND AGREES THAT ALL DOUBTS SHALL BE CONSTRUED IN FAVOR OF THE RELEASED PARTY."
    )

    # Signature Blocks
    doc.add_paragraph("\n\nBy: ______________________________")
    doc.add_paragraph("Grantor")

    doc.add_paragraph("\nBy: ______________________________")
    doc.add_paragraph("Grantee")

    doc.add_paragraph(f"\nWitness: ____________________________\n{data['witness']}")

    output_path = f"Personal_Property_Release_{data['grantor'].replace(' ', '_')}.docx"
    doc.save(output_path)
    return output_path
