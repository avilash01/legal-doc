from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import uuid
from datetime import datetime

import firebase_admin
from firebase_admin import credentials, firestore

app = Flask(__name__)

appId = globals().get('__app_id', 'default-app-id')
firebase_config_str = globals().get('__firebase_config', '{}')
initialAuthToken = globals().get('__initial_auth_token', None)

firebaseConfig = {}
try:
    if firebase_config_str:
        firebaseConfig = json.loads(firebase_config_str)
except json.JSONDecodeError:
    firebaseConfig = {}

db = None
try:
    if firebaseConfig and firebaseConfig.get("type") == "service_account":
        if not firebase_admin._apps:
            cred = credentials.Certificate(firebaseConfig)
            firebase_admin.initialize_app(cred)
        db = firestore.client()
except Exception as e:
    pass

USER_ID = "demo_user_id"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    court = request.form['court']
    location = request.form['location']
    mp_no = request.form['mp_no']
    cc_no = request.form['cc_no']
    petitioner_name = request.form['petitioner_name']
    respondent_name = request.form['respondent_name']
    reason_for_absence = request.form['reason_for_absence']
    place = request.form['place']
    date = request.form['date']

    petition_data = {
        "court": court,
        "location": location,
        "mp_no": mp_no,
        "cc_no": cc_no,
        "petitioner_name": petitioner_name,
        "respondent_name": respondent_name,
        "reason_for_absence": reason_for_absence,
        "place": place,
        "date": date,
        "timestamp": firestore.SERVER_TIMESTAMP
    }

    if db:
        try:
            doc_ref = db.collection(f'artifacts/{appId}/users/{USER_ID}/petitions').add(petition_data)
        except Exception:
            pass

    document = Document()

    p = document.add_paragraph()
    run = p.add_run(f"IN THE COURT OF {court.upper()}")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"{location.upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"M. P. No.    {mp_no} of \t {date.split('-')[0]}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"In")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"C. C. No.   {cc_no}  of {date.split('-')[0]}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    p = document.add_paragraph("BETWEEN:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"{petitioner_name.ljust(80)} … PETITIONERS/ACCUSED.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"{' ' * 70}VS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f"{respondent_name.ljust(80)} ... RESPONDENT/COMPLAINANT.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    p = document.add_paragraph("PETITION FILED UNDER SEC. 311 Cr.P.C.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"The petitioner above named states as follows:")
    document.add_paragraph(f"1. The petitioner states that his counsel on record was unable to appear before this Hon’ble Court on {date}.")
    document.add_paragraph("2. The petitioner states that it is neither willful nor wanton but due to the reason stated above.")
    document.add_paragraph("3. The petitioner states that the balance of convenience lies in his favour, further to get rebuttal evidence from the PW.")
    document.add_paragraph("In these circumstances it is therefore prayed that this Hon’ble Court may be pleased to permit the petitioner to reopen and recall the PW – and cross examine in the interest of justice and thus render justice.")
    document.add_paragraph("")

    p = document.add_paragraph(f"Signature of Counsel & Deponent ")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'petition_{petitioner_name.replace(" ", "_")}_{date}.docx'
    )

@app.route('/view_petitions')
def view_petitions():
    petitions = []
    if db:
        try:
            petitions_ref = db.collection(f'artifacts/{appId}/users/{USER_ID}/petitions')
            docs = petitions_ref.stream()
            for doc in docs:
                petition_data = doc.to_dict()
                petition_data['id'] = doc.id
                if 'timestamp' in petition_data and hasattr(petition_data['timestamp'], 'strftime'):
                    petition_data['timestamp'] = petition_data['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
                elif 'timestamp' in petition_data:
                    petition_data['timestamp'] = str(petition_data['timestamp'])
                petitions.append(petition_data)
            petitions.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        except Exception:
            pass
    return render_template('past_petitions.html', petitions=petitions, user_id=USER_ID)

if __name__ == '__main__':
    app.run(debug=True)
